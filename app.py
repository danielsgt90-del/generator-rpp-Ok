import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.units import inch
import datetime
import io
import json
from openai import OpenAI # Library untuk menghubungkan ke AI

# --- KONFIGURASI HALAMAN ---
st.set_page_config(page_title="Generator RPP AI - Deep Learning", page_icon="🤖", layout="wide")

# --- FUNGSI AI (OTAK APLIKASI) ---
def generate_rpp_content(topik, kelas, fase, mapel, api_key):
    """
    Fungsi ini mengirim permintaan ke OpenAI untuk mengisi data RPP secara otomatis.
    """
    try:
        client = OpenAI(api_key=api_key)
        
        prompt = f"""
        Anda adalah ahli kurikulum merdeka Indonesia. Buatkan data RPP Deep Learning dalam format JSON untuk topik berikut:
        Topik: {topik}
        Mata Pelajaran: {mapel}
        Kelas: {kelas}
        Fase: {fase}

        Harap isi kolom berikut dalam format JSON yang valid (tanpa markdown code block, hanya JSON murni):
        {{
            "satuan_pendidikan": "Nama Sekolah Contoh",
            "nama_guru": "Nama Guru",
            "mata_pelajaran": "{mapel}",
            "kelas": "{kelas}",
            "semester": "Ganjil/Genap",
            "fase": "{fase}",
            "elemen_pokok": "{topik}",
            "alokasi_waktu": "3 x 3 JP",
            "t1_peserta_didik": "Analisis kesiapan peserta didik...",
            "t1_materi_pelajaran": "Jenis pengetahuan (Faktual/Konseptual/Prosedural), relevansi...",
            "t1_profil_lulusan": "Pilih dimensi profil pelajar pancasila yang relevan...",
            "t1_pertanyaan_pemantik": "1. Pertanyaan memantik... 2. ...",
            "t1_sarana": "Daftar sarana fisik dan digital...",
            "t2_cp": "Capaian Pembelajaran dari kurikulum merdeka...",
            "t2_tp": "Tujuan Pembelajaran spesifik menggunakan KKO (Kata Kerja Operasional)...",
            "t2_pemahaman_bermakna": "Esensi pembelajaran...",
            "t2_lintas_disiplin": "Mata pelajaran lain yang relevan...",
            "t2_topik": "Sub-topik bahasan...",
            "t2_pedagogis": "Model pembelajaran (PjBL/PBL/Inquiry)...",
            "t2_kemitraan": "Mitra pembelajaran...",
            "t2_lingkungan": "Penataan kelas dan budaya belajar...",
            "t2_digital": "Aplikasi digital yang digunakan...",
            "t3_awal": "Kegiatan pembuka (Salam, doa, apersepsi, pemantik)...",
            "t3_awal_prinsip": "Prinsip: Menggembirakan...",
            "t3_inti": "Kegiatan inti (Memahami, Mengaplikasi, Merefleksi)...",
            "t3_inti_prinsip": "Prinsip: Berkesadaran...",
            "t3_penutup": "Kegiatan penutup (Refleksi, rangkuman)...",
            "t3_penutup_prinsip": "Prinsip: Bermakna...",
            "t4_diagnostik": "Teknik asesmen diagnostik...",
            "t4_diagnostik_kriteria": "Indikator diagnostik...",
            "t4_formatif": "Teknik asesmen formatif...",
            "t4_formatif_kriteria": "Indikator formatif...",
            "t4_sumatif": "Teknik asesmen sumatif...",
            "t4_sumatif_kriteria": "Indikator sumatif...",
            "t4_tindak_lanjut": "Remedial dan pengayaan..."
        }}
        """

        response = client.chat.completions.create(
            model="gpt-3.5-turbo", # Model yang cepat dan murah, atau gpt-4o untuk lebih pintar
            messages=[
                {"role": "system", "content": "Anda adalah asisten ahli Kurikulum Merdeka Indonesia yang hanya menjawab dalam format JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7
        )
        
        content = response.choices[0].message.content
        # Bersihkan karakter aneh jika ada (kadang AI suka menambah ```json)
        if content.startswith("```json"):
            content = content[7:]
        if content.endswith("```"):
            content = content[:-3]
            
        return json.loads(content)

    except Exception as e:
        st.error(f"Terjadi kesalahan pada AI: {e}")
        return None

# --- FUNGSI MEMBUAT DOKUMEN (Sama seperti sebelumnya) ---
def set_cell_shading(cell, fill_color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_word_doc(data):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    title = doc.add_heading('RENCANA PEMBELAJARAN MENDALAM (DEEP LEARNING)', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # INFO UMUM
    info_table = doc.add_table(rows=7, cols=2)
    info_data = [
        ("SATUAN PENDIDIKAN", data.get('satuan_pendidikan', '-')),
        ("NAMA GURU", data.get('nama_guru', '-')),
        ("MATA PELAJARAN", data.get('mata_pelajaran', '-')),
        ("KELAS / SEMESTER", f"{data.get('kelas', '-')} / {data.get('semester', '-')}"),
        ("FASE", data.get('fase', '-')),
        ("ELEMEN/MATERI POKOK", data.get('elemen_pokok', '-')),
        ("ALOKASI WAKTU", data.get('alokasi_waktu', '-')),
    ]
    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i].cells
        row[0].text = label
        row[1].text = f": {value}"
        row[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()

    # Helper Tabel
    def add_word_table(doc, title, subtitle, headers, content_rows):
        doc.add_heading(title, level=1)
        if subtitle:
            p = doc.add_paragraph(subtitle)
            p.runs[0].italic = True
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            set_cell_shading(hdr_cells[i], "D9D9D9")
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        for row_data in content_rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
                if i == 0 and len(headers) == 2:
                     if row_cells[i].paragraphs[0].runs:
                        row_cells[i].paragraphs[0].runs[0].bold = True

    # TABEL 1
    t1_rows = [
        ("Peserta Didik", data.get('t1_peserta_didik', '-')),
        ("Materi Pelajaran", data.get('t1_materi_pelajaran', '-')),
        ("Dimensi Profil Lulusan", data.get('t1_profil_lulusan', '-')),
        ("Pertanyaan Pemantik", data.get('t1_pertanyaan_pemantik', '-')),
        ("Sarana & Prasarana", data.get('t1_sarana', '-')),
    ]
    add_word_table(doc, "TABEL 1: IDENTIFIKASI", "(Analisis awal)", ["Aspek", "Deskripsi Analitis"], t1_rows)

    # TABEL 2
    t2_rows = [
        ("Capaian Pembelajaran", data.get('t2_cp', '-')),
        ("Tujuan Pembelajaran", data.get('t2_tp', '-')),
        ("Pemahaman Bermakna", data.get('t2_pemahaman_bermakna', '-')),
        ("Lintas Disiplin Ilmu", data.get('t2_lintas_disiplin', '-')),
        ("Topik Pembelajaran", data.get('t2_topik', '-')),
        ("Praktik Pedagogis", data.get('t2_pedagogis', '-')),
        ("Kemitraan Pembelajaran", data.get('t2_kemitraan', '-')),
        ("Lingkungan & Budaya", data.get('t2_lingkungan', '-')),
        ("Pemanfaatan Digital", data.get('t2_digital', '-')),
    ]
    add_word_table(doc, "TABEL 2: DESAIN PEMBELAJARAN", "(Peta konsep)", ["Komponen", "Rumusan"], t2_rows)

    # TABEL 3
    t3_rows = [
        ("KEGIATAN AWAL", data.get('t3_awal', '-'), data.get('t3_awal_prinsip', '-')),
        ("KEGIATAN INTI", data.get('t3_inti', '-'), data.get('t3_inti_prinsip', '-')),
        ("KEGIATAN PENUTUP", data.get('t3_penutup', '-'), data.get('t3_penutup_prinsip', '-')),
    ]
    add_word_table(doc, "TABEL 3: PENGALAMAN BELAJAR", "(Memahami, Mengaplikasi, Merefleksi)", ["Tahap", "Kegiatan", "Prinsip"], t3_rows)

    # TABEL 4
    t4_rows = [
        ("Asesmen Diagnostik", data.get('t4_diagnostik', '-'), data.get('t4_diagnostik_kriteria', '-')),
        ("Asesmen Formatif", data.get('t4_formatif', '-'), data.get('t4_formatif_kriteria', '-')),
        ("Asesmen Sumatif", data.get('t4_sumatif', '-'), data.get('t4_sumatif_kriteria', '-')),
        ("Tindak Lanjut", data.get('t4_tindak_lanjut', '-'), "Sesuai hasil asesmen"),
    ]
    add_word_table(doc, "TABEL 4: ASESMEN", "(Penilaian)", ["Jenis", "Teknik", "Kriteria"], t4_rows)

    # Simpan
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- TAMPILAN UTAMA (UI) ---

st.title("🤖 Generator RPP Otomatis (AI Powered)")
st.markdown("Aplikasi ini menggunakan **Kecerdasan Buatan (AI)** untuk membuat RPP secara otomatis. Anda hanya perlu memasukkan topik!")

# Sidebar untuk Konfigurasi Rahasia
with st.sidebar:
    st.header("🔑 Pengaturan AI")
    api_key_input = st.text_input("Masukkan OpenAI API Key", type="password", help="Kunci ini tidak disimpan server, hanya digunakan sesaat untuk generate.")
    st.markdown("---")
    st.info("💡 Tip: Jika tidak punya API Key, aplikasi akan mengisi data manual.")

# Form Input Sederhana
st.header("📝 Input Data Dasar")
col1, col2 = st.columns(2)
with col1:
    mapel = st.text_input("Mata Pelajaran", "Pendidikan Agama Kristen")
    kelas = st.selectbox("Kelas", ["X", "XI", "XII"])
with col2:
    fase = st.selectbox("Fase", ["E", "F", "G"])
    topik = st.text_input("Topik / Materi Pokok", placeholder="Contoh: Gereja dan Masyarakat Majemuk")

generate_button = st.button("✨ Generate RPP dengan AI", type="primary")

# Logika Saat Tombol Ditekan
if generate_button:
    if not topik:
        st.warning("Mohon isi Topik/Materi Pokok terlebih dahulu.")
    elif not api_key_input:
        st.error("Mohon masukkan OpenAI API Key di sidebar kiri untuk mengaktifkan AI.")
    else:
        with st.spinner("🧠 AI sedang menyusun RPP... Mohon tunggu sekitar 10-20 detik..."):
            # 1. Panggil AI untuk mendapatkan data
            data_ai = generate_rpp_content(topik, kelas, fase, mapel, api_key_input)
            
            if data_ai:
                st.success("✅ RPP berhasil dibuat oleh AI!")
                st.session_state['rpp_data'] = data_ai # Simpan data di session
                
                # Tampilkan preview
                with st.expander("👀 Lihat Hasil AI (Preview)"):
                    st.json(data_ai)

if 'rpp_data' in st.session_state:
    data_final = st.session_state['rpp_data']
    st.header("📥 Download Dokumen")
    col_btn1, col_btn2, col_btn3 = st.columns(3)
    
    # Generate Word
    word_buffer = create_word_doc(data_final)
    with col_btn1:
        st.download_button(
            label="📄 Download Word (.docx)",
            data=word_buffer,
            file_name=f"RPP_{data_final.get('elemen_pokok', 'AI')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    with col_btn2:
        st.markdown("🛈 Format PDF & Google Docs bisa ditambahkan kode yang sama seperti sebelumnya.")
